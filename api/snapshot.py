import asyncio
import os
import tempfile
import httpx
import json
import random
import hashlib
from bs4 import BeautifulSoup
from dotenv import load_dotenv
from playwright.async_api import async_playwright
from langchain_openai import ChatOpenAI
from langchain_core.prompts import ChatPromptTemplate
from html_parser import extract_readable_page
from typing import Literal, List, Dict, Tuple
from docx import Document
from docx.shared import Inches, Pt

# ReportLab imports
from reportlab.lib.pagesizes import A4, letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, 
    PageBreak, Image, Frame, PageTemplate, KeepTogether
)
from reportlab.lib.units import inch, mm
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
from reportlab.pdfgen import canvas
from reportlab.platypus.tableofcontents import TableOfContents

# WeasyPrint imports
from weasyprint import HTML, CSS

load_dotenv()

# ==========================================
# PREMIUM COLOR SCHEMES & DESIGN SYSTEMS
# ==========================================

PREMIUM_COLOR_SCHEMES = {
    "midnight_blue": {
        "primary": "#1E3A8A",      # Deep blue
        "secondary": "#3B82F6",     # Bright blue
        "accent": "#60A5FA",        # Light blue
        "text": "#1E293B",          # Dark slate
        "light_bg": "#F1F5F9",      # Slate 100
        "name": "Midnight Executive"
    },
    "emerald_luxe": {
        "primary": "#047857",       # Emerald
        "secondary": "#10B981",     # Green
        "accent": "#34D399",        # Light green
        "text": "#064E3B",          # Dark green
        "light_bg": "#ECFDF5",      # Green 50
        "name": "Emerald Luxe"
    },
    "royal_purple": {
        "primary": "#6D28D9",       # Purple
        "secondary": "#8B5CF6",     # Violet
        "accent": "#A78BFA",        # Light purple
        "text": "#3730A3",          # Indigo
        "light_bg": "#F5F3FF",      # Purple 50
        "name": "Royal Purple"
    },
    "crimson_edge": {
        "primary": "#BE123C",       # Rose
        "secondary": "#E11D48",     # Pink red
        "accent": "#FB7185",        # Light rose
        "text": "#881337",          # Dark rose
        "light_bg": "#FFF1F2",      # Rose 50
        "name": "Crimson Edge"
    },
    "sunset_gold": {
        "primary": "#D97706",       # Amber
        "secondary": "#F59E0B",     # Yellow
        "accent": "#FBBF24",        # Light amber
        "text": "#78350F",          # Dark amber
        "light_bg": "#FFFBEB",      # Amber 50
        "name": "Sunset Gold"
    },
    "slate_professional": {
        "primary": "#475569",       # Slate
        "secondary": "#64748B",     # Light slate
        "accent": "#94A3B8",        # Lighter slate
        "text": "#1E293B",          # Dark slate
        "light_bg": "#F8FAFC",      # Slate 50
        "name": "Slate Professional"
    },
    "teal_modern": {
        "primary": "#0F766E",       # Teal
        "secondary": "#14B8A6",     # Bright teal
        "accent": "#2DD4BF",        # Light teal
        "text": "#134E4A",          # Dark teal
        "light_bg": "#F0FDFA",      # Teal 50
        "name": "Teal Modern"
    },
    "indigo_tech": {
        "primary": "#4338CA",       # Indigo
        "secondary": "#6366F1",     # Light indigo
        "accent": "#818CF8",        # Lighter indigo
        "text": "#312E81",          # Dark indigo
        "light_bg": "#EEF2FF",      # Indigo 50
        "name": "Indigo Tech"
    },
    "rose_elegant": {
        "primary": "#9F1239",       # Deep rose
        "secondary": "#BE123C",     # Rose
        "accent": "#FB7185",        # Pink
        "text": "#4C0519",          # Darkest rose
        "light_bg": "#FFF1F2",      # Rose 50
        "name": "Rose Elegant"
    },
    "ocean_depth": {
        "primary": "#0E7490",       # Cyan
        "secondary": "#06B6D4",     # Bright cyan
        "accent": "#22D3EE",        # Light cyan
        "text": "#164E63",          # Dark cyan
        "light_bg": "#ECFEFF",      # Cyan 50
        "name": "Ocean Depth"
    }
}

LAYOUT_STYLES = {
    "corporate": {
        "cover_layout": "centered",
        "section_style": "traditional",
        "accent_position": "top",
        "title_size": 48,
        "use_boxes": True,
        "divider_style": "line"
    },
    "modern": {
        "cover_layout": "left_aligned",
        "section_style": "bold_headers",
        "accent_position": "left",
        "title_size": 52,
        "use_boxes": True,
        "divider_style": "thick_accent"
    },
    "minimalist": {
        "cover_layout": "centered",
        "section_style": "clean",
        "accent_position": "subtle",
        "title_size": 44,
        "use_boxes": False,
        "divider_style": "thin_line"
    },
    "bold": {
        "cover_layout": "left_aligned",
        "section_style": "statement",
        "accent_position": "full_width",
        "title_size": 56,
        "use_boxes": True,
        "divider_style": "gradient_bar"
    },
    "elegant": {
        "cover_layout": "centered",
        "section_style": "refined",
        "accent_position": "top_bottom",
        "title_size": 46,
        "use_boxes": True,
        "divider_style": "double_line"
    }
}

def get_design_system(url: str) -> Tuple[dict, dict]:
    """
    Generate a consistent but unique design system based on URL hash.
    Same URL will always get the same design, different URLs get different designs.
    """
    # Create hash from URL for consistency
    url_hash = hashlib.md5(url.encode()).hexdigest()
    
    # Use hash to deterministically select color scheme and layout
    color_schemes = list(PREMIUM_COLOR_SCHEMES.keys())
    layout_styles = list(LAYOUT_STYLES.keys())
    
    # Convert hash to index
    color_index = int(url_hash[:8], 16) % len(color_schemes)
    layout_index = int(url_hash[8:16], 16) % len(layout_styles)
    
    selected_color = PREMIUM_COLOR_SCHEMES[color_schemes[color_index]]
    selected_layout = LAYOUT_STYLES[layout_styles[layout_index]]
    
    return selected_color, selected_layout

load_dotenv()

# Initialize LLM
llm = ChatOpenAI(
    model="gpt-4o-mini",
    temperature=0.8,
    api_key=os.getenv("OPENAI_API_KEY")
)

class DynamicCanvas(canvas.Canvas):
    """Custom canvas with dynamic color schemes for headers and footers"""
    def __init__(self, *args, **kwargs):
        self.template_type = kwargs.pop('template_type', 'marketing')
        self.color_scheme = kwargs.pop('color_scheme', PREMIUM_COLOR_SCHEMES['midnight_blue'])
        self.layout_style = kwargs.pop('layout_style', LAYOUT_STYLES['modern'])
        canvas.Canvas.__init__(self, *args, **kwargs)
        self.pages = []
        
    def showPage(self):
        self.pages.append(dict(self.__dict__))
        self._startPage()
        
    def save(self):
        page_count = len(self.pages)
        for page_num, page_dict in enumerate(self.pages, 1):
            self.__dict__.update(page_dict)
            if page_num > 1:  # Skip header/footer on cover page
                self.draw_page_decorations(page_num, page_count)
            canvas.Canvas.showPage(self)
        canvas.Canvas.save(self)
        
    def draw_page_decorations(self, page_num, page_count):
        """Draw dynamic headers and footers based on layout style"""
        primary_color = colors.HexColor(self.color_scheme['primary'])
        accent_color = colors.HexColor(self.color_scheme['accent'])
        
        if self.layout_style['accent_position'] == 'top':
            # Top accent line
            self.setStrokeColor(primary_color)
            self.setLineWidth(3)
            self.line(30*mm, 280*mm, 180*mm, 280*mm)
            
        elif self.layout_style['accent_position'] == 'left':
            # Left vertical accent
            self.setStrokeColor(primary_color)
            self.setLineWidth(4)
            self.line(25*mm, 40*mm, 25*mm, 280*mm)
            
        elif self.layout_style['accent_position'] == 'full_width':
            # Full width top bar
            self.setFillColor(primary_color)
            self.rect(0, 285*mm, 210*mm, 5*mm, fill=1, stroke=0)
            
        elif self.layout_style['accent_position'] == 'top_bottom':
            # Top and bottom lines
            self.setStrokeColor(primary_color)
            self.setLineWidth(2)
            self.line(30*mm, 280*mm, 180*mm, 280*mm)
            self.setStrokeColor(accent_color)
            self.setLineWidth(1)
            self.line(30*mm, 20*mm, 180*mm, 20*mm)
        
        # Page number - styled based on layout
        self.setFont('Helvetica', 9)
        self.setFillColor(colors.HexColor(self.color_scheme['text']))
        
        if self.layout_style['cover_layout'] == 'left_aligned':
            self.drawString(30*mm, 15*mm, f"{page_num}")
        else:
            self.drawRightString(180*mm, 15*mm, f"{page_num}")
        
        # Footer accent - subtle
        self.setStrokeColor(colors.HexColor("#DFE6E9"))
        self.setLineWidth(0.5)
        self.line(30*mm, 20*mm, 180*mm, 20*mm)

def create_dynamic_styles(color_scheme: dict, layout_style: dict):
    """Create a premium, dynamic style palette based on color scheme and layout"""
    styles = getSampleStyleSheet()
    
    # Extract colors
    primary = colors.HexColor(color_scheme['primary'])
    secondary = colors.HexColor(color_scheme['secondary'])
    accent = colors.HexColor(color_scheme['accent'])
    text_color = colors.HexColor(color_scheme['text'])
    light_bg = colors.HexColor(color_scheme['light_bg'])
    
    # Determine alignment based on layout
    cover_align = TA_CENTER if layout_style['cover_layout'] == 'centered' else TA_LEFT
    
    # Cover page title - dynamic and impactful
    styles.add(ParagraphStyle(
        name='DynamicCoverTitle',
        parent=styles['Heading1'],
        fontSize=layout_style['title_size'],
        leading=layout_style['title_size'] + 6,
        textColor=text_color,
        spaceAfter=12,
        alignment=cover_align,
        fontName='Helvetica-Bold',
        leftIndent=0,
    ))
    
    # Cover subtitle
    styles.add(ParagraphStyle(
        name='DynamicCoverSubtitle',
        parent=styles['Normal'],
        fontSize=18,
        leading=24,
        textColor=secondary,
        spaceAfter=6,
        alignment=cover_align,
        fontName='Helvetica',
    ))
    
    # Cover metadata
    styles.add(ParagraphStyle(
        name='DynamicCoverMeta',
        parent=styles['Normal'],
        fontSize=11,
        leading=16,
        textColor=colors.HexColor("#636E72"),
        spaceAfter=6,
        alignment=cover_align,
        fontName='Helvetica',
    ))
    
    # Section headers - bold with dynamic sizing
    section_size = 28 if layout_style['section_style'] in ['statement', 'bold_headers'] else 24
    styles.add(ParagraphStyle(
        name='DynamicSection',
        parent=styles['Heading1'],
        fontSize=section_size,
        leading=section_size + 4,
        textColor=text_color,
        spaceBefore=24,
        spaceAfter=16,
        fontName='Helvetica-Bold',
        borderPadding=8,
        leftIndent=0,
    ))
    
    # Subsection headers
    styles.add(ParagraphStyle(
        name='DynamicSubsection',
        parent=styles['Heading2'],
        fontSize=16,
        leading=20,
        textColor=secondary,
        spaceBefore=16,
        spaceAfter=10,
        fontName='Helvetica-Bold',
    ))
    
    # Body text - clean and readable
    styles.add(ParagraphStyle(
        name='DynamicBody',
        parent=styles['Normal'],
        fontSize=11,
        leading=18,
        textColor=text_color,
        spaceAfter=12,
        alignment=TA_JUSTIFY,
        fontName='Helvetica',
    ))
    
    # Highlight box text - uses primary color
    styles.add(ParagraphStyle(
        name='DynamicHighlight',
        parent=styles['Normal'],
        fontSize=12,
        leading=16,
        textColor=colors.white,
        spaceAfter=10,
        alignment=TA_LEFT,
        fontName='Helvetica-Bold',
        leftIndent=12,
        rightIndent=12,
    ))
    
    # Call-out text - uses secondary color
    styles.add(ParagraphStyle(
        name='DynamicCallOut',
        parent=styles['Normal'],
        fontSize=14,
        leading=20,
        textColor=primary,
        spaceAfter=12,
        spaceBefore=12,
        alignment=TA_LEFT,
        fontName='Helvetica-Bold',
        leftIndent=20,
        borderPadding=10,
    ))
    
    # Bullet points
    styles.add(ParagraphStyle(
        name='DynamicBullet',
        parent=styles['Normal'],
        fontSize=11,
        leading=16,
        textColor=text_color,
        spaceAfter=8,
        leftIndent=20,
        bulletIndent=10,
        fontName='Helvetica',
    ))
    
    return styles

def create_dynamic_accent_box(width, height, color_scheme: dict, alpha=0.1):
    """Create a colored accent box with dynamic colors"""
    from reportlab.platypus import Flowable
    
    class AccentBox(Flowable):
        def __init__(self, width, height, color, alpha):
            Flowable.__init__(self)
            self.width = width
            self.height = height
            self.color = colors.HexColor(color)
            self.alpha = alpha
            
        def draw(self):
            self.canv.setFillColor(self.color, alpha=self.alpha)
            self.canv.rect(0, 0, self.width, self.height, fill=1, stroke=0)
    
    return AccentBox(width, height, color_scheme['primary'], alpha)

def create_dynamic_divider(color_scheme: dict, layout_style: dict):
    """Create a dynamic section divider based on style"""
    from reportlab.platypus import Flowable
    
    class Divider(Flowable):
        def __init__(self, color, style):
            Flowable.__init__(self)
            self.width = 150*mm
            self.height = 4 if style == 'thick_accent' else 2
            self.color = colors.HexColor(color)
            self.style = style
            
        def draw(self):
            if self.style == 'double_line':
                # Double line divider
                self.canv.setStrokeColor(self.color)
                self.canv.setLineWidth(2)
                self.canv.line(0, 2, 60*mm, 2)
                self.canv.setLineWidth(1)
                self.canv.line(0, 0, 60*mm, 0)
            elif self.style == 'thick_accent':
                # Thick accent bar
                self.canv.setStrokeColor(self.color)
                self.canv.setLineWidth(4)
                self.canv.line(0, 0, 80*mm, 0)
            elif self.style == 'gradient_bar':
                # Solid colored bar (gradient simulation)
                self.canv.setFillColor(self.color)
                self.canv.rect(0, 0, 60*mm, 3, fill=1, stroke=0)
            else:
                # Thin line (default)
                self.canv.setStrokeColor(self.color)
                self.canv.setLineWidth(2)
                self.canv.line(0, 0, 60*mm, 0)
    
    divider_style = layout_style.get('divider_style', 'line')
    return Divider(color_scheme['primary'], divider_style)

def render_dynamic_premium_pdf(data: Dict, output_path: str, url: str) -> str:
    """
    Creates a stunning, ultra-premium PDF with dynamic colors and layouts.
    Each URL gets a unique but consistent design.
    """
    # Get dynamic design system
    color_scheme, layout_style = get_design_system(url)
    
    print(f"\n🎨 Design System Selected:")
    print(f"   Color Scheme: {color_scheme['name']}")
    print(f"   Layout Style: {list(LAYOUT_STYLES.keys())[list(LAYOUT_STYLES.values()).index(layout_style)]}")
    print(f"   Primary Color: {color_scheme['primary']}")
    
    # Custom page template with dynamic canvas
    doc = SimpleDocTemplate(
        output_path,
        pagesize=A4,
        rightMargin=30*mm,
        leftMargin=30*mm,
        topMargin=25*mm,
        bottomMargin=25*mm,
    )
    
    styles = create_dynamic_styles(color_scheme, layout_style)
    story = []
    
    # Extract colors for easy access
    primary = colors.HexColor(color_scheme['primary'])
    secondary = colors.HexColor(color_scheme['secondary'])
    accent = colors.HexColor(color_scheme['accent'])
    light_bg = colors.HexColor(color_scheme['light_bg'])
    
    # ==========================================
    # COVER PAGE - Dynamic and Impactful
    # ==========================================
    
    # Top spacing varies by layout
    if layout_style['cover_layout'] == 'centered':
        story.append(Spacer(1, 50*mm))
    else:
        story.append(Spacer(1, 35*mm))
    
    # Accent bar - style varies
    if layout_style['accent_position'] == 'full_width':
        story.append(create_dynamic_accent_box(150*mm, 6*mm, color_scheme, 1.0))
    elif layout_style['accent_position'] in ['top', 'top_bottom']:
        story.append(create_dynamic_accent_box(150*mm, 4*mm, color_scheme, 1.0))
    elif layout_style['accent_position'] == 'left':
        story.append(create_dynamic_accent_box(8*mm, 100*mm, color_scheme, 0.15))
        story.append(Spacer(1, -100*mm))  # Overlap
    
    story.append(Spacer(1, 10*mm))
    
    # Main title with dynamic styling
    title = data.get("title", "Intelligence Report")
    story.append(Paragraph(title, styles['DynamicCoverTitle']))
    
    # Subtitle/tagline
    subtitle = data.get("subtitle", "Comprehensive Analysis & Strategic Insights")
    story.append(Paragraph(subtitle, styles['DynamicCoverSubtitle']))
    
    story.append(Spacer(1, 8*mm))
    
    # Metadata in a clean format
    url_display = data.get("url", "N/A")
    from datetime import datetime
    date_str = datetime.now().strftime("%B %d, %Y")
    
    story.append(Paragraph(f"<b>Source:</b> {url_display}", styles['DynamicCoverMeta']))
    story.append(Paragraph(f"<b>Generated:</b> {date_str}", styles['DynamicCoverMeta']))
    story.append(Paragraph(f"<b>Design:</b> {color_scheme['name']}", styles['DynamicCoverMeta']))
    
    story.append(Spacer(1, 15*mm))
    
    # Add a highlight box with key insight - uses primary color
    if data.get("key_insight") and layout_style['use_boxes']:
        highlight_data = [[Paragraph(data["key_insight"], styles['DynamicHighlight'])]]
        highlight_table = Table(highlight_data, colWidths=[150*mm])
        highlight_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), primary),
            ('PADDING', (0, 0), (-1, -1), 15),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ]))
        story.append(highlight_table)
    elif data.get("key_insight"):
        # No box version - just styled text
        story.append(Paragraph(f"<i>{data['key_insight']}</i>", styles['DynamicCallOut']))
    
    story.append(PageBreak())
    
    # ==========================================
    # EXECUTIVE SUMMARY (if provided)
    # ==========================================
    
    if data.get("executive_summary"):
        story.append(Paragraph("Executive Summary", styles['DynamicSection']))
        story.append(create_dynamic_divider(color_scheme, layout_style))
        story.append(Spacer(1, 6*mm))
        
        # Add summary in a subtle highlight box if layout uses boxes
        if layout_style['use_boxes']:
            summary_data = [[Paragraph(data["executive_summary"], styles['DynamicBody'])]]
            summary_table = Table(summary_data, colWidths=[150*mm])
            summary_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), light_bg),
                ('PADDING', (0, 0), (-1, -1), 15),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('LEFTPADDING', (0, 0), (-1, -1), 20),
                ('RIGHTPADDING', (0, 0), (-1, -1), 20),
            ]))
            story.append(summary_table)
        else:
            # Minimalist version without box
            story.append(Paragraph(data["executive_summary"], styles['DynamicBody']))
        
        story.append(Spacer(1, 8*mm))
    
    # ==========================================
    # MAIN CONTENT SECTIONS
    # ==========================================
    
    sections = data.get("sections", [])
    for idx, section in enumerate(sections):
        # Section header with dynamic styling
        heading = section.get("heading", f"Section {idx+1}")
        story.append(Paragraph(heading, styles['DynamicSection']))
        story.append(create_dynamic_divider(color_scheme, layout_style))
        story.append(Spacer(1, 4*mm))
        
        # Section content
        content = section.get("content", "")
        
        # Handle different content types
        if isinstance(content, str):
            # Split into paragraphs for better readability
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    # Check if it's a bullet point
                    if para.strip().startswith('•') or para.strip().startswith('-'):
                        story.append(Paragraph(para.strip(), styles['DynamicBullet']))
                    else:
                        story.append(Paragraph(para.strip(), styles['DynamicBody']))
        
        elif isinstance(content, list):
            # Render as modern bullet points
            for item in content:
                bullet_text = f"• {item}"
                story.append(Paragraph(bullet_text, styles['DynamicBullet']))
        
        story.append(Spacer(1, 6*mm))
        
        # Add subsections if they exist
        if section.get("subsections"):
            for subsection in section["subsections"]:
                story.append(Paragraph(subsection.get("heading", ""), styles['DynamicSubsection']))
                story.append(Paragraph(subsection.get("content", ""), styles['DynamicBody']))
                story.append(Spacer(1, 4*mm))
    
    # ==========================================
    # KEY INSIGHTS / HIGHLIGHTS (if provided)
    # ==========================================
    
    if data.get("key_points"):
        story.append(PageBreak())
        story.append(Paragraph("Key Insights", styles['DynamicSection']))
        story.append(create_dynamic_divider(color_scheme, layout_style))
        story.append(Spacer(1, 6*mm))
        
        for point in data["key_points"]:
            story.append(Paragraph(f"• {point}", styles['DynamicCallOut']))
    
    # ==========================================
    # STRATEGIC REFERENCES (Links)
    # ==========================================
    
    links = data.get("links", [])
    if links:
        story.append(PageBreak())
        story.append(Paragraph("Strategic References", styles['DynamicSection']))
        story.append(create_dynamic_divider(color_scheme, layout_style))
        story.append(Spacer(1, 6*mm))
        
        # Create a modern table with dynamic colors
        table_data = []
        table_data.append([
            Paragraph("<b>Resource</b>", styles['DynamicBody']),
            Paragraph("<b>URL</b>", styles['DynamicBody'])
        ])
        
        for link in links[:40]:
            table_data.append([
                Paragraph(link.get("text", "N/A"), styles['DynamicBody']),
                Paragraph(link.get("href", "#"), styles['DynamicBody'])
            ])
        
        link_table = Table(table_data, colWidths=[60*mm, 90*mm])
        link_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), primary),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 11),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('TOPPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.white),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#DFE6E9")),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING', (0, 0), (-1, -1), 10),
            ('RIGHTPADDING', (0, 0), (-1, -1), 10),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, light_bg]),
        ]))
        story.append(link_table)
    
    # Build the PDF with dynamic canvas
    doc.build(
        story, 
        canvasmaker=lambda *args, **kwargs: DynamicCanvas(
            *args, 
            template_type='marketing', 
            color_scheme=color_scheme,
            layout_style=layout_style,
            **kwargs
        )
    )
    
    return output_path

async def capture_page(url: str, output_path: str, format: Literal["pdf", "png"], html: str = None):
    """Captures the website as it is in PDF or PNG format using the URL directly or provided HTML."""
    async with async_playwright() as p:
        browser = await p.chromium.launch(args=["--no-sandbox", "--disable-setuid-sandbox"])
        context = await browser.new_context(
            user_agent="Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/121.0.0.0 Safari/537.36"
        )
        page = await context.new_page()
        
        if html:
            await page.set_content(html, wait_until="networkidle")
        else:
            try:
                await page.goto(url, wait_until="networkidle", timeout=60000)
            except Exception:
                try:
                    await page.goto(url, wait_until="domcontentloaded", timeout=30000)
                except:
                    pass
        
        if format == "png":
            # 🆕 Rolling screenshot logic: scroll to bottom slowly
            await page.evaluate("""
                async () => {
                    await new Promise((resolve) => {
                        let totalHeight = 0;
                        let distance = 100;
                        let timer = setInterval(() => {
                            let scrollHeight = document.body.scrollHeight;
                            window.scrollBy(0, distance);
                            totalHeight += distance;
                            if(totalHeight >= scrollHeight){
                                clearInterval(timer);
                                resolve();
                            }
                        }, 100);
                    });
                }
            """)
            await page.wait_for_timeout(1000) # Wait for final renders
            await page.screenshot(path=output_path, full_page=True)
        elif format == "pdf":
            if html:
                HTML(string=html).write_pdf(output_path)
            else:
                await page.pdf(path=output_path, format="A4", print_background=True)
            
        await browser.close()

async def scrape_site_data(url: str, html: str = None) -> dict:
    """
    Scrapes the website using BeautifulSoup to extract key components.
    """
    if not html:
        headers = {
            "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/121.0.0.0 Safari/537.36",
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,image/apng,*/*;q=0.8",
            "Accept-Language": "en-US,en;q=0.9",
        }
        
        async with httpx.AsyncClient(timeout=30.0, follow_redirects=True, headers=headers) as client:
            try:
                resp = await client.get(url)
                resp.raise_for_status()
                html = resp.text
            except Exception as e:
                print(f"BeautifulSoup Scrape Error for {url}: {repr(e)}")
                return {
                    "title": "Website Data", 
                    "links": [], 
                    "images": [], 
                    "textContent": f"Note: Could not fetch live content from URL ({str(e)}).", 
                    "html": ""
                }

    soup = BeautifulSoup(html, "html.parser")
    
    # Extract links
    links = []
    for a in soup.find_all("a", href=True):
        text = a.get_text(strip=True)
        if text:
            links.append({"text": text, "href": a["href"]})
            
    # Extract images
    images = []
    for img in soup.find_all("img", src=True):
        alt = img.get("alt") or img.get("title") or ""
        images.append({"alt": alt, "src": img["src"]})
        
    readable = extract_readable_page(html)
    textContent = readable.get("content", soup.get_text(separator=" ", strip=True))

    return {
        "title": soup.title.string.strip() if soup.title else "Website Snapshot",
        "links": links[:100],
        "images": images[:100],
        "textContent": textContent,
        "html": html
    }

async def html_to_pdf(html_content: str, output_path: str):
    """Converts raw HTML content to a PDF file using WeasyPrint."""
    try:
        HTML(string=html_content).write_pdf(output_path)
    except Exception as e:
        print(f"WeasyPrint Error: {e}")
        HTML(string=f"<h1>Rendering Error</h1><p>{str(e)}</p>").write_pdf(output_path)

async def generate_markdown_report(url: str, html: str = None) -> str:
    """Scrapes the site and returns the content in Markdown format."""
    dom_data = await scrape_site_data(url, html=html)
    title = dom_data.get("title", "Untitled Page")
    content = dom_data.get("textContent", "")
    return f"# {title}\n\nSource: {url}\n\n{content}"

def get_markdown(url: str, html: str) -> str:
    """Extracts readable content from provided HTML and returns it in Markdown format."""
    parsed = extract_readable_page(html)
    title = parsed.get("head", {}).get("title", "Untitled Page")
    content = parsed.get("content", "")
    return f"# {title}\n\nSource: {url}\n\n{content}"

async def generate_smart_pdf(url: str, template: str = "marketing", html: str = None, progress_callback=None) -> str:
    """
    Generates a premium, modern PDF report using enhanced ReportLab templates.
    """
    # Scrape the data
    if progress_callback: await progress_callback(10, "Scraping website data...")
    dom_data = await scrape_site_data(url, html=html)
    
    title = dom_data.get("title", "Website Report")
    links = dom_data.get("links", [])
    images = dom_data.get("images", [])
    text = dom_data.get("textContent", "")[:40000]

    # Enhanced template configurations
    templates = {
        "marketing": {
            "role": "Senior Marketing Strategist and Brand Consultant",
            "tone": "Professional, engaging, and action-oriented with data-driven insights",
            "instructions": """Create a comprehensive marketing intelligence report with:
            
            1. An executive summary highlighting key opportunities and threats
            2. Market positioning analysis with competitive insights
            3. Customer journey touchpoints and conversion opportunities
            4. Content strategy recommendations
            5. Key performance indicators and success metrics
            
            Structure the content into clear, actionable sections. Each section should provide strategic value.
            Include specific examples, data points, and actionable recommendations wherever possible.
            
            Focus on:
            - Brand messaging and value proposition
            - Target audience insights
            - Competitive advantages
            - Growth opportunities
            - Marketing channels and tactics"""
        },
        "business": {
            "role": "Management Consultant and Business Analyst",
            "tone": "Analytical, precise, and strategic",
            "instructions": """Create a strategic business analysis with:
            
            1. Executive summary with key findings and recommendations
            2. Business model analysis
            3. Market opportunity assessment
            4. Operational insights and efficiency opportunities
            5. Strategic recommendations and next steps
            
            Maintain a professional, data-driven approach throughout."""
        },
        "smart": {
            "role": "Product Strategy Director and UX Researcher",
            "tone": "Modern, user-centric, and innovation-focused",
            "instructions": """Create a product intelligence report with:
            
            1. Executive overview of product/service positioning
            2. User experience analysis and insights
            3. Feature analysis and competitive comparison
            4. Innovation opportunities
            5. Strategic product recommendations"""
        }
    }

    t = templates.get(template, templates["marketing"])

    # Enhanced prompt for better structured output
    prompt = ChatPromptTemplate.from_template("""
You are a {role}. Analyze the following website content and create a comprehensive, professionally structured report.

Website Title: {title}
URL: {url}
Content Summary: {text}

CRITICAL INSTRUCTIONS:
1. Return ONLY valid JSON - no markdown code blocks, no explanations
2. Create a report that is visually scannable and professionally formatted
3. Use clear, concise language that delivers immediate value
4. Structure content hierarchically with main sections and subsections where appropriate

Required JSON Structure:
{{
    "title": "Compelling Report Title (50-80 characters)",
    "subtitle": "Engaging subtitle that summarizes value (80-120 characters)",
    "key_insight": "Single most important insight in 1-2 sentences (for cover page highlight)",
    "executive_summary": "Comprehensive 3-4 paragraph executive summary covering main findings and recommendations",
    "key_points": [
        "Key insight 1",
        "Key insight 2",
        "Key insight 3",
        "Key insight 4",
        "Key insight 5"
    ],
    "sections": [
        {{
            "heading": "Section Title",
            "content": "Detailed analysis with multiple paragraphs separated by \\n\\n. Include specific examples, data points, and insights from the source content. Make this substantive and valuable.",
            "subsections": [
                {{
                    "heading": "Subsection Title",
                    "content": "Supporting details and analysis"
                }}
            ]
        }}
    ],
    "links": [
        {{"text": "Resource name", "href": "URL"}},
        {{"text": "Another resource", "href": "URL"}}
    ]
}}

Tone: {tone}

Specific Instructions: {instructions}

IMPORTANT: 
- Make the executive summary compelling and actionable
- Each section should be 3-5 paragraphs minimum with deep insights
- Include 4-6 main sections with subsections where appropriate
- Extract 5-7 key insights for the key_points array
- Ensure all content is based on the actual website content provided
""")
    
    if progress_callback: await progress_callback(30, "Analyzing content with AI...")
    chain = prompt | llm
    response = await chain.ainvoke({
        "role": t["role"],
        "title": title,
        "url": url,
        "text": text,
        "tone": t["tone"],
        "instructions": t["instructions"]
    })
    
    if progress_callback: await progress_callback(70, "Parsing AI response...")
    try:
        # Extract JSON from response
        output_str = response.content.strip()
        if "```json" in output_str:
            output_str = output_str.split("```json")[1].split("```")[0].strip()
        elif "```" in output_str:
            output_str = output_str.split("```")[1].split("```")[0].strip()
            
        data = json.loads(output_str)
    except Exception as e:
        print(f"JSON Parse Error: {e}")
        # Fallback structure
        data = {
            "title": title,
            "subtitle": "Comprehensive Analysis & Strategic Insights",
            "key_insight": "This report provides a detailed analysis of the website's content, structure, and strategic positioning.",
            "executive_summary": text[:2000],
            "key_points": [
                "Detailed content analysis",
                "Strategic insights",
                "Market positioning",
                "Competitive advantages",
                "Growth opportunities"
            ],
            "sections": [
                {
                    "heading": "Content Analysis",
                    "content": text[:5000]
                }
            ],
            "links": links[:30]
        }

    # Ensure links are included
    if not data.get("links"):
        data["links"] = links[:30]

    # Create output file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
    output_path = temp_file.name
    temp_file.close()

    # Generate the dynamic premium PDF
    if progress_callback: await progress_callback(90, "Rendering PDF layout...")
    return render_dynamic_premium_pdf(data, output_path, url)

async def generate_word_doc(url: str, html: str = None) -> str:
    """Generates a structured MS Word (.docx) document."""
    dom_data = await scrape_site_data(url, html=html)
    
    doc = Document()
    
    title = dom_data.get("title", "Website Snapshot")
    doc.add_heading(title, 0)
    
    p = doc.add_paragraph('Source: ')
    p.add_run(url).italic = True
    
    doc.add_heading('Comprehensive Content Export', level=1)
    textContent = dom_data.get("textContent", "")
    
    paragraphs = textContent.split('\n\n')
    for para in paragraphs:
        if para.strip():
            doc.add_paragraph(para.strip())
    
    links = dom_data.get("links", [])
    if links:
        doc.add_heading('Structural Reference: Links', level=1)
        table = doc.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Label'
        hdr_cells[1].text = 'Destination URL'
        for link in links[:50]:
            row_cells = table.add_row().cells
            row_cells[0].text = link.get('text', 'N/A')
            row_cells[1].text = link.get('href', '#')

    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    output_path = temp_file.name
    doc.save(output_path)
    return output_path

async def generate_special_format(url: str, target_format: Literal["research_paper", "ppt"], html: str = None) -> str:
    """Generates Research Paper or PPT styled PDF using modern ReportLab templates."""
    dom_data = await scrape_site_data(url, html=html)
    
    title = dom_data.get("title", "Website Snapshot")
    links = dom_data.get("links", [])
    text = dom_data.get("textContent", "")[:40000]
    
    if target_format == "research_paper":
        prompt = ChatPromptTemplate.from_template("""
You are an Academic Research Editor. Create a professional research paper structure.

RETURN ONLY VALID JSON.

Data: {text}
URL: {url}

Structure:
{{
    "title": "Academic Paper Title",
    "subtitle": "Research Domain and Focus",
    "key_insight": "Primary research contribution",
    "executive_summary": "Abstract (150-250 words)",
    "sections": [
        {{ "heading": "Introduction", "content": "Background and research context..." }},
        {{ "heading": "Literature Review", "content": "Related work and theoretical framework..." }},
        {{ "heading": "Methodology", "content": "Research approach and methods..." }},
        {{ "heading": "Analysis", "content": "Findings and discussion..." }},
        {{ "heading": "Conclusions", "content": "Summary and implications..." }}
    ],
    "key_points": ["Key finding 1", "Key finding 2", "..."],
    "links": [{{"text": "Reference", "href": "URL"}}, ...]
}}
""")
    else:
        prompt = ChatPromptTemplate.from_template("""
You are a Presentation Designer. Create a slide deck structure.

RETURN ONLY VALID JSON.

Data: {text}

Structure:
{{
    "title": "Presentation Title",
    "subtitle": "Compelling tagline",
    "key_insight": "Main takeaway",
    "executive_summary": "Overview slide content",
    "sections": [
        {{ "heading": "Slide 1: Problem Statement", "content": "..." }},
        {{ "heading": "Slide 2: Market Opportunity", "content": "..." }},
        ...
    ],
    "key_points": ["Point 1", "Point 2", "..."]
}}
""")
        
    chain = prompt | llm
    response = await chain.ainvoke({"title": title, "url": url, "text": text})
    
    try:
        output_str = response.content.strip()
        if "```json" in output_str:
            output_str = output_str.split("```json")[1].split("```")[0].strip()
        elif "```" in output_str:
            output_str = output_str.split("```")[1].split("```")[0].strip()
        data = json.loads(output_str)
    except:
        data = {
            "title": f"{target_format.replace('_', ' ').title()} - {title}",
            "subtitle": "Analysis Report",
            "key_insight": "Comprehensive analysis of web content",
            "executive_summary": text[:1000],
            "sections": [{"heading": "Content", "content": text[:10000]}],
            "key_points": ["Insight 1", "Insight 2"],
            "links": links[:20]
        }
    
    if not data.get("links"):
        data["links"] = links[:20]
        
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
    output_path = temp_file.name
    temp_file.close()
    
    return render_dynamic_premium_pdf(data, output_path, url)